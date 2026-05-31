"""
Styling engine for the Sanad research paper.
Centralises the teal clinical palette, fonts, heading styles, table theming,
figure/caption helpers, footer page numbers and TOC field insertion so the
content module (content.py) can stay declarative.
"""
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand palette (sampled from the live Sanad Tailwind UI) -----------------
TEAL_700   = RGBColor(0x0F, 0x76, 0x6E)   # headings
TEAL_600   = RGBColor(0x0D, 0x94, 0x88)   # primary accent
TEAL_800   = RGBColor(0x11, 0x55, 0x4E)   # deep accent
TEAL_950   = RGBColor(0x04, 0x2F, 0x2E)   # cover background
TEAL_050   = "F0FDFA"                      # light fill (hex string for shading)
TEAL_100   = "CCFBF1"
CYAN_700   = RGBColor(0x0E, 0x74, 0x90)
SLATE_700  = RGBColor(0x33, 0x41, 0x55)   # body text
SLATE_500  = RGBColor(0x64, 0x74, 0x8B)   # captions / secondary
ROSE_600   = RGBColor(0xE1, 0x1D, 0x48)   # crisis / danger
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "0F766E"                      # table header row fill (teal-700)
BAND_FILL   = "F0FDFA"                      # zebra band fill (teal-50)

BODY_FONT = "Calibri"
HEAD_FONT = "Calibri"


def _shade(cell, fill_hex):
    """Apply a solid background fill to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _bottom_border(paragraph, color_hex="0D9488", size=12):
    """Draw a coloured rule beneath a heading paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def setup_styles(doc):
    """Configure Normal + heading styles for the whole document."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = SLATE_700
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.25
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    specs = {
        "Heading 1": (16, TEAL_700, True, 18, 8),
        "Heading 2": (13, TEAL_800, True, 12, 4),
        "Heading 3": (11.5, CYAN_700, True, 8, 2),
    }
    for name, (size, color, bold, before, after) in specs.items():
        st = doc.styles[name]
        st.font.name = HEAD_FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(doc, text, level=1, rule=False):
    h = doc.add_heading(text, level=level)
    if rule and level == 1:
        _bottom_border(h, "0D9488", 14)
    return h


def add_body(doc, text, align="justify"):
    p = doc.add_paragraph(text)
    if align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = SLATE_500
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    return p


def add_figure(doc, image_path, caption, width_in=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    add_caption(doc, caption)


def add_table(doc, headers, rows, caption=None, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # header row
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        _shade(hdr[i], HEADER_FILL)
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(2)
        run = para.add_run(htext)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
    # body rows with zebra banding
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, val in enumerate(row):
            if r_idx % 2 == 1:
                _shade(cells[c_idx], BAND_FILL)
            para = cells[c_idx].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            run = para.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    if caption:
        add_caption(doc, caption)
    return table


def add_page_field(paragraph):
    """Insert a PAGE field into a paragraph run (for footers)."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldEnd)


def setup_footer(doc, label):
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label + "    |    Page ")
    r.font.size = Pt(8.5); r.font.color.rgb = SLATE_500
    add_page_field(p)


def add_toc(doc):
    """Insert a Word TOC field (updates on open with F9 / 'update fields')."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve")
    t.text = "Right-click and choose 'Update Field' to build the Table of Contents."
    placeholder.append(t)
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instr); run._r.append(sep)
    run._r.append(placeholder); run._r.append(fldEnd)


def page_break(doc):
    doc.add_page_break()
